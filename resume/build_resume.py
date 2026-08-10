"""Generate an ATS-friendly resume as .docx.

ATS parsers are simple text extractors. They routinely mangle or silently drop
content in headers/footers, text boxes, tables and multi-column layouts, and
they cannot read text inside images. So this document deliberately uses only:

* a single column of ordinary paragraphs, in reading order
* contact details in the BODY, never in a header
* standard section titles (SUMMARY, TECHNICAL SKILLS, PROJECTS, PROFESSIONAL
  EXPERIENCE, EDUCATION, PUBLICATIONS) that parsers match on
* a common font (Calibri) and plain hyphen bullets
* dates as plain text on the same line as the role

Every metric here is one that was measured in the project repository, not
estimated.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

NAME = "Sanjana Reddy"
CONTACT = "skomma18@umd.edu | 201-554-3566 | linkedin.com/in/sanjana-reddyk | College Park, MD"

SUMMARY = (
    "Data Science graduate student with 1.8 years of professional experience building production "
    "data pipelines, API integrations, and analytical systems at PricewaterhouseCoopers. Built an "
    "end-to-end MLOps platform with automated retraining, model registry, drift detection, and "
    "champion/challenger promotion using Apache Airflow, MLflow, FastAPI, and Docker. Hands-on "
    "experience fine-tuning transformer models (DeBERTa-v3, Mistral-7B/QLoRA) and building agentic "
    "RAG systems with LangGraph. Seeking Data Scientist, Machine Learning Engineer, or AI Engineer roles."
)

SKILLS = [
    ("Languages", "Python, SQL, C++"),
    (
        "Machine Learning & Deep Learning",
        "PyTorch, Scikit-Learn, LightGBM, Hugging Face Transformers, DeBERTa, LSTM, "
        "Gradient Boosting, Time Series Forecasting, Fine-Tuning, QLoRA/PEFT, Regression, "
        "Classification, Feature Engineering, Model Evaluation, Ablation Studies, "
        "Model Calibration (ECE), SHAP Explainability, MC-Dropout Uncertainty Estimation",
    ),
    (
        "MLOps & Model Operations",
        "Apache Airflow, MLflow (Experiment Tracking, Model Registry, Model Serving), "
        "Automated Retraining Pipelines, Champion/Challenger Model Promotion, Model Monitoring, "
        "Data Drift Detection (Evidently AI), Feature Stores, Train/Serve Parity, "
        "Model Versioning, A/B Model Evaluation",
    ),
    (
        "Generative AI & LLMs",
        "Retrieval-Augmented Generation (RAG), Agentic AI Systems, LangGraph, Prompt Engineering, "
        "Semantic Search, Vector Embeddings (pgvector), OpenAI API, Azure OpenAI Service, Claude, Gemini",
    ),
    (
        "Data Engineering",
        "ETL Pipelines, Workflow Orchestration (Airflow DAGs), Data Modeling, Data Integration, "
        "REST APIs, Streaming Ingestion, Data Quality Validation, Data Governance, PostgreSQL, "
        "MySQL, MSSQL, MongoDB, Apache HBase",
    ),
    (
        "Deployment & Infrastructure",
        "FastAPI, Docker, Docker Compose, Git, CI/CD (GitHub Actions), Unit & Integration Testing, "
        "Performance Profiling & Optimization, Microsoft Azure (Compute, Containers, Virtual Networks)",
    ),
    (
        "Analytics & Visualization",
        "Statistical Analysis, Data Mining, Exploratory Data Analysis, Streamlit, Matplotlib, Power BI",
    ),
]

PROJECTS = [
    (
        "Real-Time Bike-Share Demand Forecasting Platform (MLOps)",
        "GitHub: github.com/Sanjana132/Citbike_MLOps",
        [
            "Built an end-to-end MLOps platform forecasting next-hour bike demand across 2,436 NYC "
            "stations from 54.6M historical trips, orchestrating ingestion, training, monitoring, and "
            "serving across 8 containerized services (Apache Airflow, MLflow, FastAPI, PostgreSQL, "
            "Evidently AI, Streamlit, Docker Compose).",
            "Closed the retraining loop end to end: hourly monitoring of rolling MAE and feature drift "
            "automatically triggers retraining, champion/challenger evaluation, and model promotion with "
            "no human intervention; achieved test MAE of 20.27, a 56% improvement over the seasonal-naive "
            "baseline and 17.5% over the prior model.",
            "Reduced prediction API latency 13x (310 ms to 23 ms) by profiling the serving path, "
            "eliminating a per-request external API call, caching feature history, and moving prediction "
            "logging to background tasks.",
            "Diagnosed and fixed a train/serve skew in which model promotion scored an in-memory estimator "
            "while production served a different artifact, shipping a model 60% worse than the one it "
            "replaced; rebuilt promotion to score every candidate through the deployed inference path.",
            "Ran controlled feature ablations that showed engineered seasonal features degraded accuracy "
            "5.6% despite high feature importance, and removed them rather than shipping them; validated "
            "the full system with 171 hermetic automated tests, ruff linting, and GitHub Actions CI.",
        ],
    ),
    (
        "Fake News and Source Credibility Detector",
        "GitHub",
        [
            "Fine-tuned a DeBERTa-v3 regressor to score claim credibility (0-1) across 92K examples from 4 "
            "fact-checking datasets, outperforming a 0.287-MAE TF-IDF baseline.",
            "Quantified prediction uncertainty with MC-Dropout confidence intervals, measured model "
            "calibration (ECE), and generated SHAP explanations for interpretability.",
            "Built an agentic RAG layer (LangGraph orchestrating 4 retrieval tools and a QLoRA fine-tuned "
            "Mistral-7B for justification generation) and served the system through a hardened "
            "FastAPI/Docker stack with 80 automated tests and CI.",
        ],
    ),
    (
        "Txtai Market Research Platform",
        "GitHub",
        [
            "Developed a multi-agent Generative AI platform using Retrieval-Augmented Generation, semantic "
            "search, and locally hosted LLMs to automate financial due diligence.",
            "Engineered end-to-end ETL pipelines ingesting SEC filings, financial news, Reddit, and investor "
            "relations data into a vector-enabled knowledge base with pgvector embeddings, enabling document "
            "retrieval, prompt orchestration, and AI-generated research insights.",
        ],
    ),
    (
        "Car Cost Prediction",
        "GitHub",
        [
            "Built and compared regression models (Linear Regression, Decision Tree, Random Forest) to "
            "predict used car prices, with data cleaning, outlier handling, label encoding, and correlation "
            "analysis to identify key pricing drivers.",
        ],
    ),
]

EXPERIENCE = [
    (
        "PricewaterhouseCoopers - Technology Consultant",
        "August 2023 - April 2025",
        "Bangalore, India",
        [
            "Designed and implemented enterprise data capture and integration processes, producing "
            "high-quality operational and financial datasets that served as the foundation for downstream "
            "reporting and analysis.",
            "Built SQL-based data pipelines and automated data workflows that reduced manual intervention "
            "and measurably improved data accuracy across distributed retail operations.",
            "Integrated enterprise systems through real-time REST APIs, improving data availability, "
            "consistency, and automation between platforms.",
            "Developed centralized data solutions supporting 150+ retail stores and 10,000+ employees, and "
            "an expense platform processing $7.5M+ in transactions, enabling operational analytics and "
            "resource optimization at scale.",
            "Translated ambiguous business requirements into technical specifications and presented "
            "analytical findings and recommendations to non-technical stakeholders, driving data-informed "
            "decisions.",
        ],
    ),
]

EDUCATION = [
    (
        "University of Maryland, College Park - M.S., Data Science",
        "September 2025 - Present",
        "Coursework: Principles of Data Science, Principles of Machine Learning, Big Data Systems, "
        "Data Representation and Modeling",
    ),
    (
        "National Institute of Technology, Nagpur - B.Tech., Electrical and Electronics Engineering",
        "July 2019 - May 2023",
        "Coursework: Statistics and Optimization Techniques, Numerical Methods and Probability Theory, "
        "Control Systems",
    ),
]

PUBLICATIONS = [
    'K. Rajeswari, S. Neha, K. Sanjana Reddy. "Influence of AI on E-Governance and Cybersecurity in '
    'Smart Cities." International Journal of Cultural Studies and Social Sciences, Vol. XXI, No. 50, '
    "pp. 245-248."
]

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)


def _tighten(paragraph, before: int = 0, after: int = 2, spacing: float = 1.0) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = spacing


def build(path: Path) -> Path:
    document = Document()

    # One consistent, widely-installed font. Exotic fonts can break extraction.
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    for section in document.sections:
        section.top_margin = section.bottom_margin = Pt(36)
        section.left_margin = section.right_margin = Pt(40)

    # --- name and contact, in the body so parsers see them ---
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(NAME)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = ACCENT
    _tighten(heading, after=2)

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact.add_run(CONTACT)
    contact_run.font.size = Pt(9.5)
    _tighten(contact, after=6)

    def section_heading(text: str) -> None:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = ACCENT
        _tighten(paragraph, before=6, after=2)

    def bullet(text: str) -> None:
        paragraph = document.add_paragraph(text, style="List Bullet")
        paragraph.paragraph_format.left_indent = Pt(14)
        _tighten(paragraph, after=1)

    # --- summary ---
    section_heading("Summary")
    _tighten(document.add_paragraph(SUMMARY), after=4)

    # --- skills: "Label: values" lines, which parsers key on reliably ---
    section_heading("Technical Skills")
    for label, values in SKILLS:
        paragraph = document.add_paragraph()
        label_run = paragraph.add_run(f"{label}: ")
        label_run.bold = True
        paragraph.add_run(values)
        _tighten(paragraph, after=1)

    # --- projects ---
    section_heading("Projects")
    for title, link, bullets in PROJECTS:
        paragraph = document.add_paragraph()
        title_run = paragraph.add_run(title)
        title_run.bold = True
        if link:
            paragraph.add_run(f"  |  {link}").font.size = Pt(9)
        _tighten(paragraph, before=3, after=1)
        for item in bullets:
            bullet(item)

    # --- experience ---
    section_heading("Professional Experience")
    for role, dates, location, bullets in EXPERIENCE:
        paragraph = document.add_paragraph()
        role_run = paragraph.add_run(role)
        role_run.bold = True
        paragraph.add_run(f"   {dates}")
        _tighten(paragraph, before=3, after=0)

        location_paragraph = document.add_paragraph()
        location_run = location_paragraph.add_run(location)
        location_run.italic = True
        location_run.font.size = Pt(9.5)
        _tighten(location_paragraph, after=1)

        for item in bullets:
            bullet(item)

    # --- education ---
    section_heading("Education")
    for school, dates, coursework in EDUCATION:
        paragraph = document.add_paragraph()
        school_run = paragraph.add_run(school)
        school_run.bold = True
        paragraph.add_run(f"   {dates}")
        _tighten(paragraph, before=3, after=0)

        coursework_paragraph = document.add_paragraph()
        coursework_run = coursework_paragraph.add_run(coursework)
        coursework_run.italic = True
        coursework_run.font.size = Pt(9.5)
        _tighten(coursework_paragraph, after=1)

    # --- publications ---
    section_heading("Publications")
    for citation in PUBLICATIONS:
        bullet(citation)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    return path


if __name__ == "__main__":
    output = build(Path(__file__).parent / "Sanjana_Reddy_Resume_DataScience.docx")
    print(f"wrote {output}")
